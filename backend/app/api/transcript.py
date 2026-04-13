"""
POST /api/transcript/upload
===========================
Accepts a transcript file (.txt, .json, .docx) and returns structured JSON.

Auto-detects format by extension:
  .txt  → Chrome extension format (Name → HH:MM:SS → Text, blank-line separated)
  .json → Array of {speaker, timestamp, text} OR object with "turns" key
  .docx → Table-first (Speaker | Timestamp | Text columns), paragraph fallback
"""

from __future__ import annotations

import io
import json
import re
import logging
from typing import List, Optional

from fastapi import APIRouter, File, HTTPException, UploadFile
from pydantic import BaseModel

logger = logging.getLogger(__name__)
router = APIRouter()

# ── Pydantic schemas ──────────────────────────────────────────────────────────

class TranscriptTurn(BaseModel):
    speaker: str
    timestamp: str      # HH:MM:SS
    text: str


class TranscriptResponse(BaseModel):
    turns: List[TranscriptTurn]
    speakers: List[str]             # unique, in order of first appearance
    total_turns: int
    duration: str                   # timestamp of the last valid turn
    format_detected: str            # 'txt' | 'json' | 'docx_table' | 'docx_paragraphs'


# ── Shared helpers ────────────────────────────────────────────────────────────

_TIMESTAMP_RE = re.compile(r"^\d{1,2}:\d{2}:\d{2}$")


def _is_timestamp(line: str) -> bool:
    return bool(_TIMESTAMP_RE.match(line.strip()))


def _unique_ordered(items: list) -> list:
    seen: set = set()
    result = []
    for item in items:
        if item not in seen:
            seen.add(item)
            result.append(item)
    return result


def _build_response(turns: List[TranscriptTurn], fmt: str) -> TranscriptResponse:
    """Build the final response from a list of parsed turns."""
    speakers = _unique_ordered([t.speaker for t in turns])
    duration = turns[-1].timestamp if turns else "00:00:00"
    return TranscriptResponse(
        turns=turns,
        speakers=speakers,
        total_turns=len(turns),
        duration=duration,
        format_detected=fmt,
    )


# ── .txt parser (Chrome extension) ───────────────────────────────────────────

def parse_txt(raw_text: str) -> List[TranscriptTurn]:
    """
    State machine:
      0 → waiting for speaker name (non-empty, non-timestamp)
      1 → waiting for timestamp
      2 → waiting for spoken text
    Blank lines reset incomplete turns.
    """
    turns: List[TranscriptTurn] = []
    lines = [ln.rstrip() for ln in raw_text.splitlines()]

    speaker: Optional[str] = None
    timestamp: Optional[str] = None
    state = 0

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            if state != 0:
                speaker = timestamp = None
                state = 0
            continue

        if state == 0:
            if not _is_timestamp(line):
                speaker = line
                state = 1
            continue

        if state == 1:
            if _is_timestamp(line):
                timestamp = line
                state = 2
            else:
                speaker = line  # treat as new speaker name, stay in state 1
            continue

        if state == 2:
            if speaker and timestamp and line:
                turns.append(TranscriptTurn(speaker=speaker, timestamp=timestamp, text=line))
            speaker = timestamp = None
            state = 0

    return turns


# ── .json parser ──────────────────────────────────────────────────────────────

def parse_json(raw_bytes: bytes) -> List[TranscriptTurn]:
    """
    Accepts:
      1. An array of objects: [ {speaker, timestamp, text}, ... ]
      2. An object with a "turns" key: { "turns": [ ... ] }
    """
    try:
        data = json.loads(raw_bytes)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=422, detail=f"Invalid JSON: {exc}")

    items: list
    if isinstance(data, list):
        items = data
    elif isinstance(data, dict) and "turns" in data:
        items = data["turns"]
    else:
        raise HTTPException(
            status_code=422,
            detail=(
                "JSON must be an array of {speaker, timestamp, text} objects "
                "or an object with a \"turns\" key."
            ),
        )

    turns: List[TranscriptTurn] = []
    for i, item in enumerate(items):
        if not isinstance(item, dict):
            continue
        spk = str(item.get("speaker", "")).strip()
        ts  = str(item.get("timestamp", "")).strip()
        txt = str(item.get("text", "")).strip()
        if spk and ts and txt:
            turns.append(TranscriptTurn(speaker=spk, timestamp=ts, text=txt))

    return turns


# ── .docx parser ──────────────────────────────────────────────────────────────

def parse_docx(raw_bytes: bytes) -> tuple[List[TranscriptTurn], str]:
    """
    Returns (turns, sub_format).
    Tries table format first → falls back to paragraph format (same as .txt).
    """
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx is not installed on the server. Cannot process .docx files.",
        )

    doc = Document(io.BytesIO(raw_bytes))

    # ── Attempt 1: table format ──────────────────────────────────────
    turns = _parse_docx_tables(doc)
    if turns:
        return turns, "docx_table"

    # ── Attempt 2: paragraph format (Name / Timestamp / Text) ────────
    full_text = "\n".join(p.text for p in doc.paragraphs)
    turns = parse_txt(full_text)
    if turns:
        return turns, "docx_paragraphs"

    return [], "docx_paragraphs"


def _parse_docx_tables(doc) -> List[TranscriptTurn]:
    """
    Look for tables whose first row looks like [Speaker, Timestamp, Text]
    (case-insensitive, partial match). Parse subsequent rows as turns.
    """
    turns: List[TranscriptTurn] = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        # Detect header row
        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]

        # Find column indices by fuzzy match
        speaker_col = _find_col(header_cells, ("speaker", "name", "participant"))
        ts_col      = _find_col(header_cells, ("timestamp", "time", "ts"))
        text_col    = _find_col(header_cells, ("text", "content", "transcript", "spoken", "message"))

        if speaker_col is None or text_col is None:
            continue  # not the right table

        # Parse data rows
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            spk = cells[speaker_col] if speaker_col < len(cells) else ""
            ts  = cells[ts_col]      if ts_col is not None and ts_col < len(cells) else "00:00:00"
            txt = cells[text_col]    if text_col < len(cells) else ""

            if spk and txt:
                turns.append(TranscriptTurn(speaker=spk, timestamp=ts, text=txt))

    return turns


def _find_col(headers: list[str], candidates: tuple[str, ...]) -> Optional[int]:
    """Return index of the first header that contains any of the candidate strings."""
    for i, h in enumerate(headers):
        for c in candidates:
            if c in h:
                return i
    return None


# ── Endpoint ──────────────────────────────────────────────────────────────────

ALLOWED_EXTENSIONS = {".txt", ".json", ".docx"}


@router.post("/upload", response_model=TranscriptResponse)
async def upload_transcript(file: UploadFile = File(...)):
    """
    Accept a .txt / .json / .docx transcript file and return structured turns.
    Auto-detects format by extension.
    """
    filename = (file.filename or "").strip()
    ext = filename[filename.rfind("."):].lower() if "." in filename else ""

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=422,
            detail=f"Unsupported file type '{ext}'. Accepted: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    raw_bytes = await file.read()
    if not raw_bytes or not raw_bytes.strip():
        raise HTTPException(status_code=422, detail="Uploaded file is empty.")

    # ── Dispatch by extension ─────────────────────────────────────────
    turns: List[TranscriptTurn] = []
    fmt: str = ext.lstrip(".")

    if ext == ".txt":
        try:
            raw_text = raw_bytes.decode("utf-8")
        except UnicodeDecodeError:
            raw_text = raw_bytes.decode("latin-1")
        turns = parse_txt(raw_text)
        fmt = "txt"

    elif ext == ".json":
        turns = parse_json(raw_bytes)
        fmt = "json"

    elif ext == ".docx":
        turns, fmt = parse_docx(raw_bytes)

    # ── Validate result ───────────────────────────────────────────────
    if not turns:
        fmt_help = {
            ".txt":  "Expected: Speaker Name / HH:MM:SS / Text (blank line between turns).",
            ".json": 'Expected: [{speaker, timestamp, text}, ...] or {"turns": [...]}.',
            ".docx": "Expected: a table with Speaker | Timestamp | Text columns, or paragraphs in Name/Timestamp/Text format.",
        }
        raise HTTPException(
            status_code=422,
            detail=f"No valid turns parsed from {filename}. {fmt_help.get(ext, '')}",
        )

    logger.info(
        "Transcript parsed: %d turns, format=%s, file=%s",
        len(turns), fmt, filename,
    )

    return _build_response(turns, fmt)
